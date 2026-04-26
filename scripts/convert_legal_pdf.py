#!/usr/bin/env python3
"""Convert Indonesian regulation PDFs into JSON, Markdown, HTML, and DOCX."""

from __future__ import annotations

import argparse
import hashlib
import html
import json
import re
import shutil
import subprocess
import sys
import unicodedata
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt
except ImportError as exc:  # pragma: no cover - environment guard
    Document = None
    WD_ALIGN_PARAGRAPH = None
    Inches = None
    Pt = None
    DOCX_IMPORT_ERROR = exc
else:
    DOCX_IMPORT_ERROR = None


WORKING_TZ = timezone.utc

ROMAN_RE = re.compile(r"^(?=[MDCLXVI])M{0,4}(CM|CD|D?C{0,3})(XC|XL|L?X{0,3})(IX|IV|V?I{0,3})$")
PAGE_NUMBER_RE = re.compile(r"^\s*[-_–—]?\s*\d+\s*[-_–—]?\s*$")
OCR_PAGE_NUMBER_RE = re.compile(r"^\s*[-_–—]\s*[0-9tloIr!]+\s*[-_–—]\s*$", re.IGNORECASE)
ELLIPSIS_POINTER_RE = re.compile(r"(?:\.\s*){2,}$")
LAW_NUMBER_RE = re.compile(r"\bNOMOR\s+([A-Z0-9./-]+)\s+TAHUN\s+(\d{4})\b", re.IGNORECASE)
FILENAME_NUMBER_RE = re.compile(r"\b(?:NOMOR|NO\.?)\s+([A-Z0-9./-]+)\s+TAHUN\s+(\d{4})\b", re.IGNORECASE)
SK_NUMBER_RE = re.compile(r"^SK\s*No\.?\s*[0-9Il|l'MABTt\s]+\s*[ABM]?$", re.IGNORECASE)
PAGE_HEADER_RE = re.compile(
    r"^(?:PRES\s*[I!|1]?\s*D(?:E|3)?N|PRESIDEN|FRESIDEN|PRESTDEN|PRESIDE\]N|BLIK\s+INDONESIA|INDONESIA|TIEPUBLIK\s+INDONESIA|R\.?E?P[UO]BLIK\s+INDONESI[\\\/A!]*|REFUBLIK\s+INDONESIA|REPUBUK\s+INDONESIA)$",
    re.IGNORECASE,
)
ARTICLE_HEADING_RE = re.compile(r"^Pasal[\s,.;:]*([0-9OoIiLl|T\s]+[A-Z]?|[IVXLCDM]+)\s*[,.;:]?$", re.IGNORECASE)
OCR_YEAR_RE = re.compile(r"\b([12Zz][0-9OoIiLl|TtGgZz]{3})\b")


@dataclass
class Paragraph:
    id: str
    kind: str
    text: str
    part: str


def run_command(args: list[str]) -> str:
    try:
        completed = subprocess.run(args, check=True, text=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    except FileNotFoundError as exc:
        raise RuntimeError(f"Command not found: {args[0]}") from exc
    except subprocess.CalledProcessError as exc:
        raise RuntimeError(exc.stderr.strip() or f"Command failed: {' '.join(args)}") from exc
    return completed.stdout


def pdf_text(pdf_path: Path) -> str:
    return run_command(["pdftotext", "-layout", "-enc", "UTF-8", str(pdf_path), "-"])


def ocr_pdf_text(pdf_path: Path, cache_base: Path, dpi: int) -> str:
    if not shutil.which("pdftoppm"):
        raise RuntimeError("pdftoppm is required for OCR extraction but is not available in PATH.")
    if not shutil.which("tesseract"):
        raise RuntimeError("tesseract is required for OCR extraction but is not available in PATH.")

    source_hash = sha256_file(pdf_path)[:16]
    cache_dir = cache_base / f"{slugify(pdf_path.stem)}-{source_hash}-{dpi}dpi"
    page_dir = cache_dir / "pages"
    text_path = cache_dir / "ocr.txt"
    if text_path.exists():
        return text_path.read_text(encoding="utf-8")

    info = pdf_info(pdf_path)
    pages_value = info.get("Pages", "")
    if not pages_value.isdigit():
        raise RuntimeError("OCR extraction requires a detected PDF page count.")

    page_dir.mkdir(parents=True, exist_ok=True)
    page_count = int(pages_value)
    page_texts: list[str] = []
    for index in range(1, page_count + 1):
        stem = f"page-{index:04d}"
        page_txt = page_dir / f"{stem}.txt"
        if not page_txt.exists():
            if index == 1 or index == page_count or index % 25 == 0:
                print(f"OCR {pdf_path.name}: page {index}/{page_count}", file=sys.stderr)
            image_prefix = page_dir / stem
            image = page_dir / f"{stem}.tif"
            run_command([
                "pdftoppm",
                "-f",
                str(index),
                "-l",
                str(index),
                "-r",
                str(dpi),
                "-tiff",
                "-singlefile",
                str(pdf_path),
                str(image_prefix),
            ])
            completed = subprocess.run(
                ["tesseract", image.name, image.stem, "-l", "ind+eng", "--psm", "4"],
                cwd=page_dir,
                text=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
            )
            if completed.returncode != 0:
                raise RuntimeError(completed.stderr.strip() or f"OCR failed for {image.name}")
            image.unlink(missing_ok=True)
        page_texts.append(page_txt.read_text(encoding="utf-8"))

    text = "\f".join(page_texts)
    text_path.write_text(text, encoding="utf-8")
    return text


def pdf_info(pdf_path: Path) -> dict[str, str]:
    info: dict[str, str] = {}
    try:
        raw = run_command(["pdfinfo", str(pdf_path)])
    except RuntimeError:
        return info
    for line in raw.splitlines():
        if ":" not in line:
            continue
        key, value = line.split(":", 1)
        info[key.strip()] = value.strip()
    return info


def normalize_space(value: str) -> str:
    value = unicodedata.normalize("NFKC", value)
    value = value.replace("\u00a0", " ")
    value = re.sub(r"[ \t]+", " ", value)
    value = re.sub(r"\s+([,.;:])", r"\1", value)
    return value.strip()


def split_label_line(line: str) -> list[str]:
    match = re.match(r"^(Menimbang|Mengingat|Memperhatikan|Menetapkan)\s*[:;]?\s+(.+)$", line, re.IGNORECASE)
    if not match:
        return [line]
    label = match.group(1)
    rest = match.group(2).strip()
    if not rest:
        return [f"{label}:"]
    return [f"{label}:", rest]


def normalize_list_marker_line(line: str) -> str:
    line = re.sub(r"^([a-z])\s{2,}(.+)$", r"\1. \2", line)
    line = re.sub(r"^([a-z])$", r"\1.", line)
    line = re.sub(r"^(\d+[A-Z]?)\s{2,}(.+)$", r"\1. \2", line)
    line = re.sub(r"^(\d+[A-Z]?)\.\s*(?=[A-Z])", r"\1. ", line)
    return line


def normalize_article_heading(text: str) -> str:
    match = ARTICLE_HEADING_RE.match(text.strip())
    if not match:
        return text
    token = re.sub(r"\s+", "", match.group(1).strip())
    if not token:
        return text
    upper_token = token.upper()
    if re.fullmatch(r"[IVXLCDM]+", upper_token) and not re.search(r"\d", token):
        return f"Pasal {upper_token}"
    normalized = (
        token.replace("O", "0")
        .replace("o", "0")
        .replace("I", "1")
        .replace("i", "1")
        .replace("L", "1")
        .replace("l", "1")
        .replace("|", "1")
        .replace("T", "7")
    ).upper()
    if re.fullmatch(r"\d+[A-Z]?", normalized):
        number_match = re.match(r"\d+", normalized)
        if number_match and int(number_match.group(0)) == 0:
            return text
        return f"Pasal {normalized}"
    return text


def normalize_ocr_years(text: str) -> str:
    def replace_year(match: re.Match[str]) -> str:
        token = match.group(1)
        normalized = token.translate(
            str.maketrans({
                "O": "0",
                "o": "0",
                "I": "1",
                "i": "1",
                "L": "1",
                "l": "1",
                "|": "1",
                "T": "1",
                "t": "1",
                "G": "9",
                "g": "9",
                "Z": "2",
                "z": "2",
            })
        )
        return normalized if normalized.isdigit() else token

    return OCR_YEAR_RE.sub(replace_year, text)


def vehicle_token_replacement(match: re.Match[str]) -> str:
    raw = match.group(0)
    letters = re.sub(r"[^A-Za-z]", "", raw)
    uppercase_count = sum(1 for char in letters if char.isupper())
    if letters and uppercase_count >= max(1, len(letters) - 1):
        return "VEHICLE)"
    if raw[:1].isupper():
        return "Vehicle)"
    return "vehicle)"


def is_article_heading(text: str) -> bool:
    normalized = normalize_article_heading(text)
    if normalized != text:
        return True
    match = ARTICLE_HEADING_RE.match(text.strip())
    if not match:
        return False
    token = re.sub(r"\s+", "", match.group(1).strip())
    if not token:
        return False
    upper_token = token.upper()
    if re.fullmatch(r"[IVXLCDM]+", upper_token) and not re.search(r"\d", token):
        return True
    numeric_token = (
        token.replace("O", "0")
        .replace("o", "0")
        .replace("I", "1")
        .replace("i", "1")
        .replace("L", "1")
        .replace("l", "1")
        .replace("|", "1")
        .replace("T", "7")
    ).upper()
    if not re.fullmatch(r"\d+[A-Z]?", numeric_token):
        return False
    number_match = re.match(r"\d+", numeric_token)
    return bool(number_match and int(number_match.group(0)) > 0)


def split_structural_line(line: str) -> list[str]:
    normalized_heading = normalize_article_heading(line)
    if normalized_heading != line:
        return [normalized_heading]
    patterns = [
        (re.compile(r"^(Pasal[\s,.;:]*(?:\d+[A-Z]?|[IVXLCDM]+))\s+(.+)$", re.IGNORECASE), ("ayat ", "huruf ", "angka ")),
        (re.compile(r"^(Ayat\s+\(\d+[A-Z]?\))\s+(.+)$"), ()),
        (re.compile(r"^(Huruf\s+[a-z])\s+(.+)$"), ()),
        (re.compile(r"^(Angka\s+\d+[A-Z]?)\s+(.+)$"), ()),
    ]
    for pattern, inline_prefixes in patterns:
        match = pattern.match(line)
        if not match:
            continue
        rest = match.group(2).strip()
        if any(rest.lower().startswith(prefix) for prefix in inline_prefixes):
            return [line]
        if match.group(1).lower().startswith("pasal ") and rest and rest[0].islower():
            return [line]
        return [match.group(1).strip(), rest]
    return [line]


def is_noise_line(line: str) -> bool:
    stripped = normalize_space(line.strip()).replace("…", "...")
    if not stripped:
        return False
    if stripped == "[":
        return True
    if stripped == "Mengingat":
        return True
    if re.fullmatch(r"\d+\.\s+Peraturan\.?", stripped):
        return True
    if SK_NUMBER_RE.match(stripped):
        return True
    if re.search(r"\bSK\s*No\.?\b", stripped, re.IGNORECASE):
        return True
    if PAGE_HEADER_RE.match(stripped):
        return True
    if PAGE_NUMBER_RE.match(stripped):
        return True
    if OCR_PAGE_NUMBER_RE.match(stripped):
        return True
    if re.fullmatch(r"Pasal\s+[A-Za-z0-9|IlLoOtT]{1,4}(?:\s+[A-Za-z]{1,3})?", stripped, re.IGNORECASE):
        return not is_article_heading(stripped)
    article_candidate = normalize_legal_text(stripped)
    if re.match(r"^Pasal\s+[0-9A-Z]+\s*(?:\.|\u2026|\s)*$", article_candidate, re.IGNORECASE):
        return not is_article_heading(article_candidate)
    if ELLIPSIS_POINTER_RE.search(stripped):
        return True
    return False


def is_roman(value: str) -> bool:
    return bool(ROMAN_RE.match(value))


def is_heading(text: str) -> bool:
    upper = text.upper()
    if upper in {
        "TENTANG",
        "DENGAN RAHMAT TUHAN YANG MAHA ESA",
        "MEMUTUSKAN:",
        "PENJELASAN",
        "ATAS",
        "I. UMUM",
        "II. PASAL DEMI PASAL",
    }:
        return True
    if re.match(r"^BAB\s+[A-Z0-9IVXLCDM]+$", upper):
        return True
    if re.match(r"^BAGIAN\s+(KESATU|KEDUA|KETIGA|KEEMPAT|KELIMA|KEENAM|KETUJUH|KEDELAPAN|KESEMBILAN|KESEPULUH|\w+)$", upper):
        return True
    if re.match(r"^PARAGRAF\s+\w+", upper):
        return True
    if is_article_heading(text):
        return True
    if re.match(r"^(UNDANG-UNDANG|PERATURAN PEMERINTAH|PERATURAN PRESIDEN|PERATURAN MENTERI|KEPUTUSAN MENTERI)\b", upper):
        return True
    if re.match(r"^NOMOR\s+.+\s+TAHUN\s+\d{4}$", upper):
        return True
    if upper.startswith("LEMBARAN NEGARA REPUBLIK INDONESIA"):
        return True
    if upper.startswith("TAMBAHAN LEMBARAN NEGARA REPUBLIK INDONESIA"):
        return True
    return False


def is_upper_title_line(text: str) -> bool:
    upper = text.upper()
    if len(text) < 3 or upper != text:
        return False
    if re.search(r"\d", text):
        return False
    if is_heading(text):
        return True
    words = re.findall(r"[A-Z]+", upper)
    return bool(words) and len(words) <= 12


def is_list_start(text: str) -> bool:
    if re.match(r"^\(\d+[A-Z]?\)\s+", text):
        return True
    if re.match(r"^\d+[A-Z]?\.\s+", text):
        return True
    if re.match(r"^\d+[A-Z]?\s+(?=[A-Z])", text):
        return True
    if re.match(r"^[a-z]\.\s+", text):
        return True
    if re.match(r"^Ayat\s+\(\d+[A-Z]?\)$", text):
        return True
    if re.match(r"^Huruf\s+[a-z]$", text):
        return True
    if re.match(r"^Angka\s+\d+[A-Z]?$", text):
        return True
    return False


def is_new_block(text: str) -> bool:
    if is_heading(text) or is_list_start(text) or is_upper_title_line(text):
        return True
    if re.match(r"^(Menimbang|Mengingat|Memperhatikan|Menetapkan):$", text, re.IGNORECASE):
        return True
    if re.match(r"^(Dengan Persetujuan Bersama|Disahkan di|Diundangkan di)\b", text, re.IGNORECASE):
        return True
    if re.match(r"^(PRESIDEN|MENTERI|DEWAN PERWAKILAN)\b", text, re.IGNORECASE):
        return True
    if text == "dan":
        return True
    return False


def closes_before_continuation(text: str) -> bool:
    if is_heading(text):
        return True
    if re.match(r"^(Menimbang|Mengingat|Memperhatikan|Menetapkan):$", text, re.IGNORECASE):
        return True
    if re.match(r"^Ayat\s+\(\d+[A-Z]?\)$", text):
        return True
    if re.match(r"^Huruf\s+[a-z]$", text):
        return True
    if re.match(r"^Angka\s+\d+[A-Z]?$", text):
        return True
    return False


def continues_reference(current: str, line: str) -> bool:
    if not current:
        return False
    if not re.search(r"\b(ayat|huruf|angka|pasal)$", current.rstrip(), re.IGNORECASE):
        return False
    return bool(re.match(r"^(\(\d+[A-Z]?\)|[a-z]\.|\d+[A-Z]?\.|\d+[A-Z]?)(?:\s|$)", line))


def append_line(current: str, line: str) -> str:
    if not current:
        return line
    if current.endswith("-") and line and line[0].islower():
        return current + line
    return f"{current} {line}"


def clean_and_reflow(raw_text: str) -> list[str]:
    paragraphs: list[str] = []
    current = ""

    for raw_page in raw_text.split("\f"):
        for raw_line in raw_page.splitlines():
            if is_noise_line(raw_line):
                continue

            stripped = normalize_legal_text(normalize_space(raw_line))
            if not stripped:
                if current:
                    if re.search(r"\b(ayat|huruf|angka|pasal)$", current.rstrip(), re.IGNORECASE):
                        continue
                    paragraphs.append(normalize_space(current))
                    current = ""
                continue

            stripped = normalize_list_marker_line(stripped)

            expanded_lines: list[str] = []
            for label_line in split_label_line(stripped):
                expanded_lines.extend(split_structural_line(label_line))
            for line_index, line in enumerate(expanded_lines):
                line = normalize_space(line)
                if not line:
                    continue
                force_new = len(expanded_lines) > 1 and line_index > 0
                new_block = is_new_block(line) and not continues_reference(current, line)
                close_current = closes_before_continuation(current) and not continues_reference(current, line)
                if current and (force_new or new_block or close_current):
                    paragraphs.append(normalize_space(current))
                    current = line
                else:
                    current = append_line(current, line)

    if current:
        paragraphs.append(normalize_space(current))

    cleaned: list[str] = []
    for para in paragraphs:
        para = normalize_legal_text(para)
        if para:
            for split_para in split_structural_line(para):
                split_para = normalize_space(split_para)
                if split_para:
                    cleaned.append(split_para)
    return merge_split_article_headings(cleaned)


def normalize_x_noise(text: str) -> str:
    whitelist = {"Toxicity", "dioxins", "NOx", "ex", "X", "IX", "XI", "XII", "XIII", "XIV", "XV", "XIIJ"}

    def replace_word(match: re.Match[str]) -> str:
        word = match.group(0)
        if word in whitelist or word.upper() in whitelist or len(word) <= 2:
            return word
        if word.isupper():
            return word
        return word.replace("x", "k").replace("X", "K")

    return re.sub(r"\b[A-Za-z]*[xX][A-Za-z]*\b", replace_word, text)


def normalize_legal_text(text: str) -> str:
    replacements = {
        "UNDANG–UNDANG": "UNDANG-UNDANG",
        "Undang–Undang": "Undang-Undang",
        "undang–undang": "undang-undang",
        "PERATURAN–PEMERINTAH": "PERATURAN PEMERINTAH",
        "PERLI N DU NGAN": "PERLINDUNGAN",
        "PERLINDI.INGAN": "PERLINDUNGAN",
        "PENGELO I-AAN": "PENGELOLAAN",
        "TENTANC}": "TENTANG",
        "Pasa!": "Pasal",
        "Pasai": "Pasal",
        "PasaI": "Pasal",
        "MEN,IUTUSI(AN": "MEMUTUSKAN",
        "UI\\.{UM": "UMUM",
        "ENERGl": "ENERGI",
        "Keija": "Kerja",
        "keija": "kerja",
        "MAFIA ESA": "MAHA ESA",
        "be1anja": "belanja",
        "bersurnber": "bersumber",
        "rnenghasilkan": "menghasilkan",
        "ekonorni": "ekonomi",
        "rates juta": "ratus juta",
        "melaiui": "melalui",
        "Leuel": "Level",
        "Leuelatau": "Level atau",
        "Levelatau": "Level atau",
        "kebdakan": "kebijakan",
        "EksPor": "Ekspor",
        "ImPor": "Impor",
        "Pertzinan": "Perizinan",
        "Perrzinan": "Perizinan",
        "permohonanPerizinan": "permohonan Perizinan",
        "voiume": "volume",
        "periindungan": "perlindungan",
        "Ra)ryat": "Rakyat",
        "hgpermarket": "hypermarket",
        "Penyelen ggaraar;": "Penyelenggaraan",
        "Promo si": "Promosi",
        "rnemenuhi": "memenuhi",
        "memiiiki": "memiliki",
        "Aiat": "Alat",
        "perundan g-undangan": "perundang-undangan",
        "daiam": "dalam",
        "kementLrian": "kementerian",
        "rlegara": "negara",
        "l.embaran": "Lembaran",
        "Tarnbahan": "Tambahan",
        "Seruice Leuel Agreement": "Service Level Agreement",
        "Seruice Level Agreement": "Service Level Agreement",
        "fianji layanan)": "(janji layanan)",
        "iebagaimana": "sebagaimana",
        "hasi!": "hasil",
        "Hasi!": "Hasil",
        "hasil!": "hasil",
        "menghasi!kan": "menghasilkan",
        "Penghasi!": "Penghasil",
        "penghasi!": "penghasil",
        "memenuh!": "memenuhi",
        "operasiona!": "operasional",
        "profi!": "profil",
        "nila!": "nilai",
        "ha!": "hal",
        "ha! ": "hal ",
        "ftuph!": "(tujuh)",
        "dimaxsud": "dimaksud",
        "dimaxksud": "dimaksud",
        "cimaxsud": "dimaksud",
        "melakuxan": "melakukan",
        "melaxsanakan": "melaksanakan",
        "diexspor": "diekspor",
        "untux": "untuk",
        "tidax": "tidak",
        "jarax": "jarak",
        "Radionuxfida": "Radionuklida",
        "nelaxsanaan": "pelaksanaan",
        "dimanfactxan": "dimanfaatkan",
        "h,rsat": "Pusat",
        "H,rsat": "Pusat",
        "Fusat": "Pusat",
        "Pemerintah hrsat": "Pemerintah Pusat",
        "xelas": "kelas",
        "xewenangannya": "kewenangannya",
        "Sarxsi": "Sanksi",
        "Penganghuten": "Pengangkutan",
        "diakukean": "dilakukan",
        "menggunaran": "menggunakan",
        "Pengargkut": "Pengangkut",
        "dilengkani": "dilengkapi",
        "dilengxap": "dilengkap",
        "memenubi": "memenuhi",
        "Lunbah": "Limbah",
        "Limban": "Limbah",
        "Linibhah": "Limbah",
        "Lirubah": "Limbah",
        "Lirnbah": "Limbah",
        "Uraban": "Limbah",
        "PRES IOEN": "PRESIDEN",
        "PRES I DEN": "PRESIDEN",
        "PRES IDEN": "PRESIDEN",
        "PRES !DEN": "PRESIDEN",
        "PRES tDEN": "PRESIDEN",
        "PRESTDEN": "PRESIDEN",
        "PRESIDE]N": "PRESIDEN",
        "FRESIDEN": "PRESIDEN",
        "REPUBLTK": "REPUBLIK",
        "REPUBUK": "REPUBLIK",
        "REFUBLIK": "REPUBLIK",
        "R.EPUBLIK": "REPUBLIK",
        "Menirnbang": "Menimbang",
        "Nornor": "Nomor",
        "Pcrarturan": "Peraturan",
        "Pcraturan": "Peraturan",
        "Perubdhan": "Perubahan",
        "Repubik": "Republik",
        "menghasilka.n": "menghasilkan",
        "arrgka": "angka",
        "ldentification": "Identification",
        "hurr.f": "huruf",
        "tan huruf": "dan huruf",
        "ten tang": "tentang",
        "!1h*\"": "bahwa",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = text.replace("“", '"').replace("”", '"').replace("„", '"').replace("’", "'")
    text = re.sub(r"^(Menimbang)\s*[\.:]\s*cL\.", r"\1: a.", text, flags=re.IGNORECASE)
    text = text.replace("lBattery", "(Battery").replace("lbattery", "(battery")
    text = re.sub(
        r"\b(?:ve,?\s?hicl(?:q|ei|el)|uehicl(?:el|l|e)?)(?:\))?",
        vehicle_token_replacement,
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(r"\b[Ee]lectic\b", lambda match: "Electric" if match.group(0)[0].isupper() else "electric", text)
    text = text.replace("{", "(").replace("}", ")")
    text = re.sub(r"\b2\)l[g9]\b", "2019", text)
    text = re.sub(r"\(!\)", "(1)", text)
    text = re.sub(r"\((\d+)!\)", r"(\1)", text)
    text = re.sub(r"\((\d+)!", r"(\1)", text)
    text = re.sub(r"\bayat\s+!([0-9])\)", r"ayat (\1)", text, flags=re.IGNORECASE)
    text = re.sub(r"\((satu|dua|tiga|empat|lima|enam|tujuh|delapan|sembilan|sepuluh)!", r"(\1)", text, flags=re.IGNORECASE)
    text = re.sub(r"\bParagraf\s+!\b", "Paragraf 1", text, flags=re.IGNORECASE)
    text = text.replace("VI!", "VII")
    text = re.sub(r"\b(?:Fasai|Fasal|Pasat|Pasai|Passl|Pasaj|Pasi)\b", "Pasal", text)
    text = re.sub(r"\b(?:Pasal!|Pasal\)|Pasa:|Pasa\?!|Fasa\)|Puasa!)\s+", "Pasal ", text)
    text = re.sub(r"^Pasal\s+(\d{2,}):$", r"Pasal \g<1>1", text, flags=re.IGNORECASE)
    text = re.sub(r"^(Pasal\s+[0-9A-Z]+)\s*(?:\.|\u2026)+$", r"\1", text, flags=re.IGNORECASE)
    text = re.sub(r"^[A-Za-z0-9]{1,3}\s+SALINAN$", "SALINAN", text)
    text = re.sub(r"^BAB\s+\|$", "BAB I", text, flags=re.IGNORECASE)
    text = re.sub(r"\bPasal\s+S5\s+ayat\b", "Pasal 5 ayat", text, flags=re.IGNORECASE)
    text = re.sub(r"\bPasal\s+5\s+ayat\s+12\)", "Pasal 5 ayat (2)", text, flags=re.IGNORECASE)
    text = re.sub(r"\{(\d+[A-Z]?)\)", r"(\1)", text)
    text = re.sub(r"\{(\d+[A-Z]?)\b", r"(\1)", text)
    text = re.sub(r"\bayat\s+\((\d)1\b(?=\s+(?:dilaksanakan|diselenggarakan|huruf))", r"ayat (\1)", text, flags=re.IGNORECASE)
    text = re.sub(r"\((\d)1\b(?=\s)", r"(\1)", text)
    text = re.sub(r"[\{(](\d+[A-Z]?)[\\lI]\)", r"(\1)", text)
    text = re.sub(r"[\{(](\d+[A-Z]?)\\", r"(\1)", text)
    text = re.sub(r"\b2[Oo0][lI|]4\b", "2014", text)
    text = re.sub(r"\b2[Oo0][lI|]8\b", "2018", text)
    text = re.sub(r"\b2[oO]2[oO]\b", "2020", text)
    text = re.sub(r"\b2[oO]2[rRIl|]\b", "2021", text)
    text = re.sub(r"\b2[OoIil|]14\b", "2014", text)
    text = re.sub(r"\b2[OoIil|]18\b", "2018", text)
    text = re.sub(r"\bZOZ[|Il1]\b", "2021", text)
    text = re.sub(r"(?<=\d)O(?=\d|\b)", "0", text)
    text = re.sub(r"\b2OOg\b", "2009", text)
    text = re.sub(r"\b2OO9\b", "2009", text)
    text = re.sub(r"\b2O2O\b", "2020", text)
    text = re.sub(r"\b2O21\b", "2021", text)
    text = normalize_ocr_years(text)
    text = re.sub(r"\bPasal\s+(\d+)l[:,]uruf\b", r"Pasal \1 huruf", text, flags=re.IGNORECASE)
    text = re.sub(r"\bPasal\s+(\d+)[lI]\b", r"Pasal \g<1>1", text, flags=re.IGNORECASE)
    text = re.sub(r"\bPasal\s+[lI](?=\d)", "Pasal 1", text)
    text = re.sub(r"\bNomor\s+S\s+Tahun\s+2021\b", "Nomor 5 Tahun 2021", text, flags=re.IGNORECASE)
    text = re.sub(r"\bNomor\s+1\s+I\s+Tahun\b", "Nomor 11 Tahun", text, flags=re.IGNORECASE)
    text = re.sub(r"\b[Pp]en5rusun", lambda match: "Penyusun" if match.group(0)[0].isupper() else "penyusun", text)
    text = re.sub(r"\b[Mm]en5rusun", lambda match: "Menyusun" if match.group(0)[0].isupper() else "menyusun", text)
    text = re.sub(r"\b[Ss]ertilikat\b", lambda match: "Sertifikat" if match.group(0)[0].isupper() else "sertifikat", text)
    text = text.replace("bersertilikat", "bersertifikat").replace("Bersertilikat", "Bersertifikat")
    text = re.sub(r"\blzin\b", "Izin", text)
    text = re.sub(r"\bdanf\s+ataulzin\b", "dan/atau Izin", text, flags=re.IGNORECASE)
    text = re.sub(r"\bd\.?anf\s+atau\b", "dan/atau", text, flags=re.IGNORECASE)
    text = re.sub(r"\b(dan)\s*(?:f|l)?\s*atau\b", r"\1/atau", text, flags=re.IGNORECASE)
    text = re.sub(r"\bdan/ataulzin\b", "dan/atau Izin", text, flags=re.IGNORECASE)
    text = re.sub(r"\bdanlatau\b", "dan/atau", text, flags=re.IGNORECASE)
    text = re.sub(r"(?<=[a-z])danlatau\b", " dan/atau", text, flags=re.IGNORECASE)
    text = re.sub(r"\bdan/ataul\b", "dan/atau", text, flags=re.IGNORECASE)
    text = re.sub(r"\bd\^\s+", "dan ", text, flags=re.IGNORECASE)
    text = re.sub(r"(?<=[a-z])lzin\b", " Izin", text)
    text = normalize_x_noise(text)
    text = text.replace("!", "l")
    text = text.replace("pemegan g Izin", "pemegang Izin")
    text = re.sub(r"\blfrequently asked questionsl\b", "(frequently asked questions)", text, flags=re.IGNORECASE)
    text = text.replace("perLlndang", "perundang")
    text = text.replace("2,5o/o", "2,5%").replace("2,5oh", "2,5%")
    text = text.replace("6Ooh", "60%")
    text = text.replace("!;ruruf", "huruf")
    text = text.replace("T\\.rgas", "Tugas")
    text = re.sub(r"\b(?:TIEPUBLIK|R\.?E?P[UO]BLIK|REFUBLIK|REPUBUK|REPUBLIK)\s+INDONES![A]?\b", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\b(?:TIEPUBLIK|R\.?E?P[UO]BLIK|REFUBLIK|REPUBUK)\s+INDONESI[\\\/A!]*\b", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\b(?:REPUtsL[IT]K|REPUtsUK|REPUEL[IT]K|REPUE[IU]K|REPUEILIK|REPUBL\|K|REPUBLIK\s+tNDONESlA|FRES\s*IDEN\s+REPUtsLIK)\s+INDONES[IASTAbl]*\b", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\bREPUBLIK\s+(?:lNDONESIA|tNDONESlA|INDONESTA|INDONEbIA)\b", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\s*\.\.\.r\.i\s+dengan ketentuan\b", " sesuai dengan ketentuan", text, flags=re.IGNORECASE)
    text = re.sub(r"\((\d)1\b(?=\s)", r"(\1)", text)
    text = re.sub(r"\((\d+)t\)", r"(\1)", text)
    text = re.sub(r"\((\d+)t\b", r"(\1)", text)
    text = re.sub(r"\bayat\s+\((\d+)1\)huruf\b", r"ayat (\1) huruf", text, flags=re.IGNORECASE)
    text = re.sub(r"\bayat\s+\((\d+)1huruf\b", r"ayat (\1) huruf", text, flags=re.IGNORECASE)
    text = re.sub(r"\bNomor\s+l1\b", "Nomor 11", text)
    text = re.sub(r"\b6573l[';]?", "6573)", text)
    text = re.sub(r"\b2I\.", "21.", text)
    text = text.replace("Pengadaan. Barang/Jasa", "Pengadaan Barang/Jasa")
    text = text.replace("Rp100.900.000,00 (seratus juta rupiah)", "Rp100.000.000,00 (seratus juta rupiah)")
    text = text.replace("penerbita n P erizinan Be ru saha secara oto mati s", "penerbitan Perizinan Berusaha secara otomatis")
    text = text.replace("pencabut an P erizinan Beru saha", "pencabutan Perizinan Berusaha")
    text = text.replace("pencabutanPerizinan", "pencabutan Perizinan")
    text = text.replace("menteri-yang -.ry\"irggarakan", "menteri yang menyelenggarakan")
    text = text.replace("y\". memasarkan", "yang memasarkan")
    text = re.sub(r"p[\".\-]*b[\".]*t[\".]*an", "pembatasan", text)
    text = text.replace("penyelen ggaraan", "penyelenggaraan")
    text = text.replace("berbasi s", "berbasis")
    text = text.replace("dan,diperlukan", "dan diperlukan")
    text = text.replace("mernenuh i", "memenuhi")
    text = text.replace("p.-..u,t d\"g^t g", "pameran dagang")
    text = text.replace("i'asal ll7", "Pasal 117")
    text = text.replace("p\"1\"k*^rr\"rtt", "pelaksanaan")
    text = text.replace("tupati/wali kota", "bupati/wali kota")
    text = text.replace("kabupate n I kota", "kabupaten/kota")
    text = text.replace("kabupatenlkotayang", "kabupaten/kota yang")
    text = text.replace("kabupatenlkota", "kabupaten/kota")
    text = text.replace("kabupaten/ kota", "kabupaten/kota")
    text = text.replace("Menteri'", "Menteri.")
    text = text.replace("administratif'", "administratif.")
    text = text.replace("Perdagangan'", "Perdagangan.")
    text = text.replace("Ayat - (", "Ayat (")
    text = text.replace("Ayat (i)", "Ayat (1)")
    text = text.replace("Ayat (a)", "Ayat (4)")
    text = text.replace("Ig4S", "1945")
    text = text.replace("Ind6nesia", "Indonesia")
    text = text.replace("menterdaskan", "mencerdaskan")
    text = text.replace("wegara Republit<", "Negara Republik")
    text = text.replace("metakukan", "melakukan")
    text = text.replace("tttndotong", "mendorong")
    text = text.replace("iu3ran", "tujuan")
    text = text.replace("warga,rf\".\" blrhak", "warga negara berhak")
    text = text.replace("keiranusiaan", "kemanusiaan")
    text = text.replace("prrgfriarp\".t y\".r[ layat", "penghidupan yang layak")
    text = text.replace("p..rIitrg", "penting")
    text = text.replace("p\"rrrU\"\"grnan", "pembangunan")
    text = re.sub(r"(?<=[A-Za-z])\.(?=[A-Za-z])", "", text)
    text = re.sub(r"(?:'\s*){2,}\s*sK\s*No\s+[A-Za-z0-9]+\s*[A-Z]?", "", text)
    text = re.sub(r"\bSK\s*No\.?\s*[0-9Il|l'MABTt]+(?:\s*[0-9Il|l'MABTt]+)*\s*[ABM]?\b", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\s+[A-Za-z]{0,4}\s+x[.)]?\s*[^A-Za-z0-9]*(?:[A-Za-z]{0,4}[.)]?)?,?\s*$", "", text)
    text = re.sub(r"\s+[-_–—]\s*[0-9tloIr!]+\s*[-_–—]\s*$", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\b([A-Za-z]+)-\s+([A-Za-z]+)\b", r"\1-\2", text)
    text = re.sub(r"\bj\s+alan\b", "jalan", text, flags=re.IGNORECASE)
    text = re.sub(r"\bperundangundangan\b", "perundang-undangan", text, flags=re.IGNORECASE)
    text = re.sub(r"\bundangundang\b", "undang-undang", text, flags=re.IGNORECASE)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def merge_split_article_headings(paragraphs: list[str]) -> list[str]:
    merged: list[str] = []
    index = 0
    while index < len(paragraphs):
        current = normalize_article_heading(paragraphs[index])
        next_text = paragraphs[index + 1] if index + 1 < len(paragraphs) else ""
        if re.match(r"^Pasal\s+\d+$", current) and re.fullmatch(r"[0-9OoIl|\s]+", next_text):
            combined = normalize_article_heading(f"{current} {next_text}")
            if combined != f"{current} {next_text}":
                merged.append(combined)
                index += 2
                continue
        merged.append(current)
        index += 1
    return merged


def classify(text: str, in_explanation: bool, before_opening_done: bool) -> str:
    upper = text.upper()
    if upper == "PENJELASAN" or upper in {"ATAS", "I. UMUM", "II. PASAL DEMI PASAL"}:
        return "explanation_heading"
    if re.match(r"^BAB\s+", upper):
        return "chapter"
    if re.match(r"^BAGIAN\s+", upper):
        return "part"
    if re.match(r"^PARAGRAF\s+", upper):
        return "subpart"
    if is_article_heading(text):
        return "article"
    if re.match(r"^\(\d+[A-Z]?\)\s+", text):
        return "paragraph"
    if re.match(r"^[a-z]\.\s+", text):
        return "letter"
    if re.match(r"^\d+[A-Z]?\.\s+", text):
        return "number"
    if re.match(r"^Ayat\s+\(\d+[A-Z]?\)$", text):
        return "explanation_item"
    if re.match(r"^Huruf\s+[a-z]$", text):
        return "explanation_item"
    if re.match(r"^Angka\s+\d+[A-Z]?$", text):
        return "explanation_item"
    if upper.startswith("LAMPIRAN"):
        return "attachment"
    if upper.startswith("LEMBARAN NEGARA") or upper.startswith("TAMBAHAN LEMBARAN NEGARA"):
        return "closing"
    if re.match(r"^(Disahkan|Diundangkan)\s+di\b", text, re.IGNORECASE):
        return "closing"
    if re.match(r"^(Menimbang|Mengingat|Memperhatikan):$", text, re.IGNORECASE):
        return "opening"
    if re.match(r"^(MEMUTUSKAN:|Menetapkan:)$", text, re.IGNORECASE):
        return "decision"
    if not before_opening_done and (is_heading(text) or is_upper_title_line(text) or upper == "DENGAN RAHMAT TUHAN YANG MAHA ESA"):
        return "title"
    if in_explanation:
        return "explanation_body"
    return "body"


def build_paragraphs(raw_paragraphs: list[str]) -> list[Paragraph]:
    result: list[Paragraph] = []
    in_explanation = False
    before_opening_done = True

    for text in raw_paragraphs:
        if not result:
            before_opening_done = False
        if re.match(r"^(Menimbang|Mengingat|Memperhatikan):$", text, re.IGNORECASE):
            before_opening_done = True
        if text.upper() == "PENJELASAN":
            in_explanation = True
        kind = classify(text, in_explanation, before_opening_done)
        part = "explanation" if in_explanation else "body"
        result.append(Paragraph(id=f"p-{len(result) + 1:04d}", kind=kind, text=text, part=part))
    return result


def normalized_article_key(text: str) -> str:
    normalized = normalize_article_heading(text)
    match = re.search(r"\b(\d+[A-Z]?)\b", normalized)
    if match:
        return match.group(1)
    roman_match = re.search(r"\b[IVXLCDM]+\b", normalized)
    return roman_match.group(0) if roman_match else normalized


def article_heading_number(text: str) -> str | None:
    normalized = normalize_article_heading(text)
    match = re.fullmatch(r"Pasal\s+([0-9A-Z]+)", normalized, flags=re.IGNORECASE)
    return match.group(1).upper() if match else None


def close_numeric_token(value: str, expected: int) -> bool:
    expected_text = str(expected)
    if value == expected_text:
        return False
    if value.startswith(expected_text) and len(value) > len(expected_text):
        return True
    if len(value) == len(expected_text):
        return sum(left != right for left, right in zip(value, expected_text)) <= 1
    return False


def coerce_article_number(token: str, expected: int | None) -> int | None:
    if token.isdigit():
        return int(token)
    if expected is None:
        return None

    candidates: list[str] = []
    if re.fullmatch(r"\d+[A-Z]", token):
        tail_map = {"O": "0", "C": "0", "D": "0", "Q": "0", "U": "0", "S": "5", "B": "8"}
        replacement = tail_map.get(token[-1])
        if replacement:
            candidates.append(token[:-1] + replacement)
    if re.fullmatch(r"[I1L|]+", token):
        candidates.append(token.replace("I", "1").replace("L", "1").replace("|", "1"))

    for candidate in candidates:
        if candidate.isdigit() and (int(candidate) == expected or close_numeric_token(candidate, expected)):
            return int(candidate)
    return None


def normalize_article_sequences(paragraphs: list[Paragraph]) -> list[Paragraph]:
    previous_number: dict[str, int | None] = {"body": None, "explanation": None}
    result: list[Paragraph] = []

    for paragraph in paragraphs:
        if paragraph.kind != "article":
            result.append(paragraph)
            continue

        token = article_heading_number(paragraph.text)
        previous = previous_number[paragraph.part]
        expected = previous + 1 if previous is not None else None
        current = coerce_article_number(token, expected) if token else None
        if current is not None:
            if expected is not None and current != expected:
                if current <= previous:
                    if current == previous:
                        pass
                    elif close_numeric_token(token, expected):
                        paragraph = Paragraph(id=paragraph.id, kind=paragraph.kind, text=f"Pasal {expected}", part=paragraph.part)
                        current = expected
                    else:
                        continue
                elif current - expected == 1:
                    pass
                elif close_numeric_token(token, expected) or (current > expected + 50 and len(token) > len(str(expected))):
                    paragraph = Paragraph(id=paragraph.id, kind=paragraph.kind, text=f"Pasal {expected}", part=paragraph.part)
                    current = expected
                elif current - expected > 20:
                    continue
            elif token != str(current):
                paragraph = Paragraph(id=paragraph.id, kind=paragraph.kind, text=f"Pasal {current}", part=paragraph.part)
            previous_number[paragraph.part] = current
        result.append(paragraph)

    return result


def remove_repeated_article_noise(paragraphs: list[Paragraph]) -> list[Paragraph]:
    result: list[Paragraph] = []
    seen: dict[str, set[str]] = {"body": set(), "explanation": set()}
    index = 0
    while index < len(paragraphs):
        paragraph = paragraphs[index]
        if paragraph.kind != "article":
            result.append(paragraph)
            index += 1
            continue

        key = normalized_article_key(paragraph.text)
        if key in seen[paragraph.part]:
            next_paragraph = paragraphs[index + 1] if index + 1 < len(paragraphs) else None
            if (
                next_paragraph
                and next_paragraph.part == paragraph.part
                and next_paragraph.kind in {"body", "explanation_body"}
                and next_paragraph.text[:1].islower()
            ):
                result.append(
                    Paragraph(
                        id="",
                        kind=next_paragraph.kind,
                        text=normalize_space(f"{paragraph.text} {next_paragraph.text}"),
                        part=paragraph.part,
                    )
                )
                index += 2
                continue
            index += 1
            continue

        seen[paragraph.part].add(key)
        result.append(paragraph)
        index += 1

    return [Paragraph(id=f"p-{idx + 1:04d}", kind=p.kind, text=p.text, part=p.part) for idx, p in enumerate(result)]


def extract_metadata(paragraphs: list[Paragraph], pdf_path: Path, info: dict[str, str], extraction_method: str) -> dict[str, object]:
    texts = [p.text for p in paragraphs]
    joined = "\n".join(texts[:80])
    doc_type = ""
    for text in texts[:20]:
        upper = text.upper()
        if upper.startswith("UNDANG-UNDANG"):
            doc_type = "UU"
            break
        if upper.startswith("PERATURAN PEMERINTAH"):
            doc_type = "PP"
            break
        if upper.startswith("PERATURAN PRESIDEN"):
            doc_type = "PERPRES"
            break
        if upper.startswith("PERATURAN MENTERI ENERGI DAN SUMBER DAYA MINERAL"):
            doc_type = "PERMEN ESDM"
            break
        if upper.startswith("PERATURAN MENTERI LINGKUNGAN HIDUP DAN KEHUTANAN"):
            doc_type = "PERMEN LHK"
            break
        if upper.startswith("PERATURAN MENTERI PERINDUSTRIAN"):
            doc_type = "PERMENPERIN"
            break
        if upper.startswith("PERATURAN MENTERI"):
            doc_type = "PERMEN"
            break

    number = ""
    year = ""
    header_joined = "\n".join(texts[:20])
    filename_joined = normalize_legal_text(pdf_path.stem.upper())
    match = LAW_NUMBER_RE.search(header_joined) or FILENAME_NUMBER_RE.search(filename_joined) or LAW_NUMBER_RE.search(joined)
    if match:
        number = match.group(1)
        year = match.group(2)

    title_lines: list[str] = []
    for i, text in enumerate(texts[:40]):
        if text.upper() != "TENTANG":
            continue
        for next_text in texts[i + 1 : i + 8]:
            upper = next_text.upper()
            if upper.startswith("DENGAN RAHMAT") or upper.startswith("PRESIDEN") or upper.startswith("MENTERI"):
                break
            title_lines.append(next_text)
        break
    title = " ".join(title_lines).strip()

    slug_type = {
        "PERMEN ESDM": "permen-esdm",
        "PERMEN LHK": "permen-lhk",
        "PERMENPERIN": "permenperin",
    }.get(doc_type, doc_type.lower() or "peraturan")
    slug_parts = [slug_type, "nomor", number.lower(), "tahun", year]
    slug = slugify(" ".join(part for part in slug_parts if part))

    return {
        "source_file": str(pdf_path),
        "source_sha256": sha256_file(pdf_path),
        "document_type": doc_type,
        "number": number,
        "year": year,
        "title": title,
        "slug": slug,
        "extraction_method": extraction_method,
        "pdf_info": {
            "pages": info.get("Pages", ""),
            "title": info.get("Title", ""),
            "tagged": info.get("Tagged", ""),
            "encrypted": info.get("Encrypted", ""),
            "pdf_version": info.get("PDF version", ""),
        },
        "generated_at": datetime.now(WORKING_TZ).isoformat(),
    }


def slugify(value: str) -> str:
    value = value.lower()
    value = re.sub(r"[^a-z0-9]+", "-", value)
    value = re.sub(r"-+", "-", value)
    return value.strip("-") or "peraturan"


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def is_non_text_attachment_paragraph(paragraph: Paragraph) -> bool:
    text = paragraph.text.strip()
    if not text:
        return True
    compact = re.sub(r"\s+", "", text)
    if re.search(r"(?:\.|\u2026){3,}", compact):
        return True
    if re.search(r"_{3,}", compact):
        return True
    if re.search(r"\b(?:PT|PLT|MW|MVA|kV|KV)\s*(?:\.|\u2026){2,}", text):
        return True
    if re.search(r"\bTAHUN\s*(?:\.|\u2026){2,}", text, re.IGNORECASE):
        return True
    if len(text) <= 2 and not re.search(r"[A-Za-z]", text):
        return True
    return False


def starts_attachment_section(text: str) -> bool:
    return bool(re.match(r"^[A-Z]\.\s+", text))


def starts_attachment_table(text: str) -> bool:
    table_patterns = [
        r"\bNo\.?\s+(?:Ketentuan|Persyaratan|Parameter|Jenis|Uraian|Kegiatan|Dokumen|Tahapan|Rekaman)\b",
        r"\btabel\s+berikut\b",
        r"\bTahapan\s+proses/\s*Alat\s+Uji/",
        r"\bUntuk\s+Perusahaan\s+Industri\s+Untuk\s+Produsen\b",
    ]
    return any(re.search(pattern, text, re.IGNORECASE) for pattern in table_patterns)


def filter_non_text_attachments(paragraphs: list[Paragraph]) -> tuple[list[Paragraph], int]:
    filtered: list[Paragraph] = []
    skipped = 0
    in_attachment = False
    in_attachment_table = False

    for paragraph in paragraphs:
        if paragraph.kind == "attachment":
            in_attachment = True
            in_attachment_table = False
            filtered.append(paragraph)
            continue
        if not in_attachment:
            filtered.append(paragraph)
            continue
        starts_section = starts_attachment_section(paragraph.text)
        if starts_section and in_attachment_table:
            in_attachment_table = False
        if starts_attachment_table(paragraph.text):
            in_attachment_table = True
            skipped += 1
            continue
        if in_attachment_table or is_non_text_attachment_paragraph(paragraph):
            skipped += 1
            continue
        filtered.append(paragraph)

    return [Paragraph(id=f"p-{idx + 1:04d}", kind=p.kind, text=p.text, part=p.part) for idx, p in enumerate(filtered)], skipped


def is_short_ocr_noise(paragraph: Paragraph) -> bool:
    if paragraph.kind in {"letter", "number"} and re.fullmatch(r"(?:[a-z]|\d+)\.\s*[^A-Za-z0-9]{1,3}", paragraph.text.strip()):
        return True
    if re.search(r"\bx\b", paragraph.text.strip()) and len(paragraph.text.strip()) <= 20:
        return True
    if paragraph.kind not in {"body", "explanation_body"}:
        return False
    text = paragraph.text.strip()
    if re.search(r"\bPUBLIK\s+INDONESIA\b.*\bDepu\b", text, re.IGNORECASE):
        return True
    if re.fullmatch(r"[*\s]*vanna\s+Djaman.*", text, flags=re.IGNORECASE):
        return True
    if re.fullmatch(r"sil\s+Djaman", text, flags=re.IGNORECASE):
        return True
    if text == 'EN ax "Be':
        return True
    if re.fullmatch(r"Pasal\s+[A-Za-z0-9|IlLoOtT]{1,4}(?:\s+[A-Za-z]{1,3})?", text, re.IGNORECASE) and not is_article_heading(text):
        return True
    if text in {"dan", "atau", "Umum"}:
        return False
    if len(text) <= 6:
        return True
    return False


def filter_ocr_noise(paragraphs: list[Paragraph]) -> tuple[list[Paragraph], int]:
    filtered: list[Paragraph] = []
    skipped = 0
    for paragraph in paragraphs:
        text = re.sub(r"^SALINAN\s+\w{1,3}$", "SALINAN", paragraph.text.strip())
        candidate = Paragraph(id=paragraph.id, kind=paragraph.kind, text=text, part=paragraph.part)
        if is_short_ocr_noise(candidate):
            skipped += 1
            continue
        filtered.append(candidate)
    return [Paragraph(id=f"p-{idx + 1:04d}", kind=p.kind, text=p.text, part=p.part) for idx, p in enumerate(filtered)], skipped


def quality_report(
    raw_text: str,
    paragraphs: list[Paragraph],
    skipped_attachment_paragraph_count: int,
    skipped_noise_paragraph_count: int,
) -> dict[str, object]:
    texts = [p.text for p in paragraphs]
    flags: list[str] = []
    joined = "\n".join(texts)
    if "\ufffd" in raw_text or "\ufffd" in joined:
        flags.append("replacement_character_found")
    if re.search(r"(?:\.\s*){3,}", joined):
        flags.append("ellipsis_pointer_possible_residue")
    if skipped_attachment_paragraph_count:
        flags.append("skipped_non_text_attachment")
    if not any(p.kind == "article" for p in paragraphs):
        flags.append("article_headings_not_detected")

    def article_number(value: str) -> str:
        normalized = normalize_article_heading(value)
        match = re.search(r"\d+[A-Z]?", normalized)
        if match:
            return match.group(0)
        roman_match = re.search(r"\b[IVXLCDM]+\b", normalized)
        return roman_match.group(0) if roman_match else normalized

    body_articles = {article_number(p.text) for p in paragraphs if p.kind == "article" and p.part == "body"}
    explanation_articles = {article_number(p.text) for p in paragraphs if p.kind == "article" and p.part == "explanation"}
    has_explanation = any(p.text.upper() == "PENJELASAN" for p in paragraphs)
    if has_explanation and body_articles and explanation_articles and len(body_articles) != len(explanation_articles):
        flags.append("body_explanation_article_count_mismatch")

    def sequence_gap_count(part: str) -> int:
        has_roman_amendment_articles = any(
            p.kind == "article"
            and p.part == part
            and re.fullmatch(r"Pasal\s+[IVXLCDM]+", normalize_article_heading(p.text), flags=re.IGNORECASE)
            for p in paragraphs
        )
        if has_roman_amendment_articles:
            return 0

        previous: int | None = None
        gaps = 0
        for paragraph in paragraphs:
            if paragraph.kind != "article" or paragraph.part != part:
                continue
            token = article_heading_number(paragraph.text)
            if not token or not token.isdigit():
                continue
            current = int(token)
            if previous is not None and current != previous + 1:
                gaps += 1
            previous = current
        return gaps

    body_sequence_gap_count = sequence_gap_count("body")
    explanation_sequence_gap_count = sequence_gap_count("explanation")
    if body_sequence_gap_count:
        flags.append("body_article_sequence_gap")
    if explanation_sequence_gap_count:
        flags.append("explanation_article_sequence_gap")

    return {
        "paragraph_count": len(paragraphs),
        "body_article_count": len(body_articles),
        "explanation_article_count": len(explanation_articles),
        "body_article_sequence_gap_count": body_sequence_gap_count,
        "explanation_article_sequence_gap_count": explanation_sequence_gap_count,
        "chapter_count": sum(1 for p in paragraphs if p.kind == "chapter"),
        "part_count": sum(1 for p in paragraphs if p.kind == "part"),
        "letter_count": sum(1 for p in paragraphs if p.kind == "letter"),
        "number_count": sum(1 for p in paragraphs if p.kind == "number"),
        "has_explanation": has_explanation,
        "has_state_gazette": "LEMBARAN NEGARA REPUBLIK INDONESIA" in joined,
        "has_supplement": "TAMBAHAN LEMBARAN NEGARA REPUBLIK INDONESIA" in joined,
        "skipped_attachment_paragraph_count": skipped_attachment_paragraph_count,
        "skipped_noise_paragraph_count": skipped_noise_paragraph_count,
        "quality_flags": flags,
    }


def regulation_payload(pdf_path: Path, output_dir: Path, force_ocr: bool, ocr_dpi: int) -> dict[str, object]:
    info = pdf_info(pdf_path)
    extraction_method = "tesseract_ocr" if force_ocr else "pdftotext_layout"
    raw_text = ocr_pdf_text(pdf_path, output_dir / "_ocr-cache", ocr_dpi) if force_ocr else pdf_text(pdf_path)
    raw_paragraphs = clean_and_reflow(raw_text)
    paragraphs = normalize_article_sequences(build_paragraphs(raw_paragraphs))
    paragraphs = remove_repeated_article_noise(paragraphs)
    paragraphs, skipped_attachment_paragraph_count = filter_non_text_attachments(paragraphs)
    paragraphs, skipped_noise_paragraph_count = filter_ocr_noise(paragraphs)
    metadata = extract_metadata(paragraphs, pdf_path, info, extraction_method)
    quality = quality_report(raw_text, paragraphs, skipped_attachment_paragraph_count, skipped_noise_paragraph_count)
    return {
        "metadata": metadata,
        "quality": quality,
        "paragraphs": [p.__dict__ for p in paragraphs],
    }


def write_json(payload: dict[str, object], out_path: Path) -> None:
    out_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def write_markdown(payload: dict[str, object], out_path: Path) -> None:
    metadata = payload["metadata"]
    paragraphs = payload["paragraphs"]
    assert isinstance(metadata, dict)
    lines = [
        "---",
        f"source_file: {json.dumps(metadata.get('source_file', ''), ensure_ascii=False)}",
        f"document_type: {json.dumps(metadata.get('document_type', ''), ensure_ascii=False)}",
        f"number: {json.dumps(metadata.get('number', ''), ensure_ascii=False)}",
        f"year: {json.dumps(metadata.get('year', ''), ensure_ascii=False)}",
        f"title: {json.dumps(metadata.get('title', ''), ensure_ascii=False)}",
        f"slug: {json.dumps(metadata.get('slug', ''), ensure_ascii=False)}",
        "---",
        "",
    ]
    for item in paragraphs:
        assert isinstance(item, dict)
        text = str(item["text"])
        kind = str(item["kind"])
        if kind == "chapter":
            lines.extend([f"## {text}", ""])
        elif kind == "article":
            lines.extend([f"### {text}", ""])
        elif kind in {"explanation_heading"} and text.upper() == "PENJELASAN":
            lines.extend([f"## {text}", ""])
        else:
            lines.extend([text, ""])
    out_path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")


def paragraph_html(text: str, kind: str, pid: str) -> str:
    escaped = html.escape(text)
    label_match = re.match(r"^((?:\(\d+[A-Z]?\)|\d+[A-Z]?\.|[a-z]\.))\s+(.+)$", text)
    if label_match and kind in {"paragraph", "number", "letter"}:
        label = html.escape(label_match.group(1))
        body = html.escape(label_match.group(2))
        inner = f'<span class="legal-label">{label}</span><span class="legal-body">{body}</span>'
    else:
        inner = escaped
    return f'<p id="{pid}" class="legal-p legal-{kind}" data-kind="{kind}">{inner}</p>'


def write_html(payload: dict[str, object], out_path: Path) -> None:
    metadata = payload["metadata"]
    paragraphs = payload["paragraphs"]
    quality = payload["quality"]
    assert isinstance(metadata, dict)
    assert isinstance(quality, dict)
    title = display_title(metadata)
    body = "\n".join(paragraph_html(str(p["text"]), str(p["kind"]), str(p["id"])) for p in paragraphs if isinstance(p, dict))
    doc = f"""<!doctype html>
<html lang="id">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{html.escape(title)}</title>
  <style>
    :root {{
      color-scheme: light;
      --paper: #ffffff;
      --ink: #222222;
      --muted: #646464;
      --rule: #dedede;
      --accent: #1f5f5b;
    }}
    body {{
      margin: 0;
      background: #f3f4f3;
      color: var(--ink);
      font-family: Arial, Helvetica, sans-serif;
      font-size: 11pt;
      line-height: 1.2;
    }}
    .document-shell {{
      max-width: 8.27in;
      margin: 32px auto;
      padding: 1in;
      background: var(--paper);
      box-shadow: 0 1px 8px rgba(0, 0, 0, 0.08);
    }}
    .document-meta {{
      border-bottom: 1px solid var(--rule);
      margin-bottom: 28px;
      padding-bottom: 14px;
      color: var(--muted);
      font-size: 10pt;
      line-height: 1.35;
    }}
    .document-meta strong {{
      color: var(--accent);
      font-weight: 700;
    }}
    .legal-p {{
      margin: 0 0 10px 0;
      text-align: justify;
      text-justify: inter-word;
    }}
    .legal-title,
    .legal-chapter,
    .legal-part,
    .legal-subpart,
    .legal-article,
    .legal-explanation_heading,
    .legal-decision,
    .legal-opening,
    .legal-closing {{
      text-align: left;
    }}
    .legal-chapter,
    .legal-explanation_heading {{
      margin-top: 18px;
    }}
    .legal-paragraph,
    .legal-letter,
    .legal-number {{
      display: grid;
      grid-template-columns: 0.42in 1fr;
      column-gap: 0.02in;
    }}
    .legal-label {{
      display: inline-block;
    }}
    .legal-body {{
      display: inline-block;
      text-align: justify;
    }}
    @media (max-width: 760px) {{
      body {{ background: var(--paper); }}
      .document-shell {{
        margin: 0;
        padding: 24px 18px;
        box-shadow: none;
      }}
      .legal-paragraph,
      .legal-letter,
      .legal-number {{
        grid-template-columns: 34px 1fr;
      }}
    }}
    @media print {{
      body {{ background: var(--paper); }}
      .document-shell {{
        margin: 0;
        box-shadow: none;
      }}
    }}
  </style>
</head>
<body>
  <main class="document-shell">
    <section class="document-meta">
      <strong>{html.escape(title)}</strong><br>
      Source: {html.escape(str(metadata.get("source_file", "")))}<br>
      Paragraphs: {html.escape(str(quality.get("paragraph_count", "")))} · Body articles: {html.escape(str(quality.get("body_article_count", "")))} · Explanation articles: {html.escape(str(quality.get("explanation_article_count", "")))}
    </section>
    <article class="legal-document">
{body}
    </article>
  </main>
</body>
</html>
"""
    out_path.write_text(doc, encoding="utf-8")


def display_title(metadata: dict[str, object]) -> str:
    doc_type = str(metadata.get("document_type") or "Peraturan")
    number = str(metadata.get("number") or "")
    year = str(metadata.get("year") or "")
    title = str(metadata.get("title") or "")
    base = " ".join(part for part in [doc_type, "Nomor", number, "Tahun", year] if part)
    if title:
        return f"{base} tentang {title}" if base else title
    return base or "Peraturan"


def write_docx(payload: dict[str, object], out_path: Path) -> None:
    if Document is None:
        raise RuntimeError(f"python-docx is unavailable: {DOCX_IMPORT_ERROR}")
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1

    for item in payload["paragraphs"]:
        assert isinstance(item, dict)
        text = str(item["text"])
        kind = str(item["kind"])
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(8)
        if kind in {"paragraph", "letter", "number"}:
            paragraph.paragraph_format.left_indent = Inches(0.42)
            paragraph.paragraph_format.first_line_indent = Inches(-0.33)
        run = paragraph.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(11)
    document.save(out_path)


def ensure_dirs(base: Path) -> dict[str, Path]:
    dirs = {
        "json": base / "json",
        "markdown": base / "markdown",
        "html": base / "html",
        "docx": base / "docx",
    }
    for path in dirs.values():
        path.mkdir(parents=True, exist_ok=True)
    return dirs


def output_stem(payload: dict[str, object], fallback: Path) -> str:
    metadata = payload["metadata"]
    assert isinstance(metadata, dict)
    slug = str(metadata.get("slug") or "")
    if slug and slug != "peraturan":
        return slug
    return slugify(fallback.stem)


def convert_pdf(pdf_path: Path, output_dir: Path, force_ocr: bool, ocr_dpi: int) -> dict[str, object]:
    payload = regulation_payload(pdf_path, output_dir, force_ocr, ocr_dpi)
    dirs = ensure_dirs(output_dir)
    stem = output_stem(payload, pdf_path)
    json_path = dirs["json"] / f"{stem}.json"
    md_path = dirs["markdown"] / f"{stem}.md"
    html_path = dirs["html"] / f"{stem}.html"
    docx_path = dirs["docx"] / f"{stem}.docx"
    write_json(payload, json_path)
    write_markdown(payload, md_path)
    write_html(payload, html_path)
    write_docx(payload, docx_path)
    payload["outputs"] = {
        "json": str(json_path),
        "markdown": str(md_path),
        "html": str(html_path),
        "docx": str(docx_path),
    }
    return payload


def iter_input_pdfs(args: argparse.Namespace) -> Iterable[Path]:
    if args.input_dir:
        yield from sorted(Path(args.input_dir).expanduser().resolve().glob("*.pdf"))
    for value in args.pdfs:
        yield Path(value).expanduser().resolve()


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("pdfs", nargs="*", help="PDF file(s) to convert.")
    parser.add_argument("--input-dir", help="Directory containing PDF files.")
    parser.add_argument("--output-dir", default="generated", help="Output directory.")
    parser.add_argument("--force-ocr", action="store_true", help="Use Tesseract OCR instead of the embedded text layer.")
    parser.add_argument("--ocr-dpi", type=int, default=300, help="Rasterization DPI for --force-ocr.")
    args = parser.parse_args(argv)

    if not shutil.which("pdftotext"):
        print("ERROR: pdftotext is required but not available in PATH.", file=sys.stderr)
        return 2

    output_dir = Path(args.output_dir).expanduser().resolve()
    output_dir.mkdir(parents=True, exist_ok=True)

    pdfs = list(dict.fromkeys(iter_input_pdfs(args)))
    if not pdfs:
        print("ERROR: provide at least one PDF or --input-dir.", file=sys.stderr)
        return 2

    index: list[dict[str, object]] = []
    failures: list[dict[str, str]] = []

    for pdf_path in pdfs:
        if not pdf_path.exists():
            failures.append({"source_file": str(pdf_path), "error": "file_not_found"})
            continue
        try:
            payload = convert_pdf(pdf_path, output_dir, args.force_ocr, args.ocr_dpi)
        except Exception as exc:  # noqa: BLE001 - CLI should continue batch conversion
            failures.append({"source_file": str(pdf_path), "error": str(exc)})
            continue
        metadata = payload["metadata"]
        quality = payload["quality"]
        outputs = payload["outputs"]
        assert isinstance(metadata, dict)
        assert isinstance(quality, dict)
        assert isinstance(outputs, dict)
        index.append({
            "source_file": str(pdf_path),
            "document_type": metadata.get("document_type"),
            "number": metadata.get("number"),
            "year": metadata.get("year"),
            "title": metadata.get("title"),
            "slug": metadata.get("slug"),
            "quality": quality,
            "outputs": outputs,
        })
        flags = ",".join(quality.get("quality_flags", [])) if isinstance(quality.get("quality_flags"), list) else ""
        print(f"OK {pdf_path.name} -> {metadata.get('slug')} flags={flags or '-'}")

    write_json({"items": index, "failures": failures}, output_dir / "index.json")
    if failures:
        for failure in failures:
            print(f"FAIL {failure['source_file']}: {failure['error']}", file=sys.stderr)
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
